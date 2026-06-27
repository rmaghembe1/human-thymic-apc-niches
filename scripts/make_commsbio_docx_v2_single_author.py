#!/usr/bin/env python3

from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_DIR = Path("/mnt/d/thymic_apc_atlas_project")
DOC_DIR = PROJECT_DIR / "docs"
PACKAGE_DIR = PROJECT_DIR / "submission_package_v3_commsbio"

IN_MD = DOC_DIR / "manuscript_draft_v6_single_author_commsbio.md"
OUT_DOCX = DOC_DIR / "manuscript_draft_v6_single_author_commsbio_numeric.docx"
OUT_MD_NUMERIC = DOC_DIR / "manuscript_draft_v6_single_author_commsbio_numeric_refs.md"
CONVERSION_REPORT = DOC_DIR / "commsbio_docx_conversion_report_v2_single_author.md"

if not IN_MD.exists():
    raise FileNotFoundError(IN_MD)

text = IN_MD.read_text()

# -----------------------------
# Parse reference keys
# -----------------------------
ref_section_match = re.search(r"^## References\n(.+?)(?=^## |\Z)", text, flags=re.S | re.M)
if not ref_section_match:
    raise ValueError("Could not find References section.")

ref_section = ref_section_match.group(1).strip()
ref_lines = [line.strip() for line in ref_section.splitlines() if line.strip()]

key_to_num = {}
for line in ref_lines:
    m = re.match(r"^(\d+)\.\s+\[([^\]]+)\]\s+", line)
    if m:
        num = int(m.group(1))
        key = m.group(2)
        key_to_num[key] = num

if not key_to_num:
    raise ValueError("No reference keys parsed.")

missing_keys = set()

def replace_citation(match):
    inside = match.group(1)
    keys = re.findall(r"@([A-Za-z0-9_]+)", inside)
    nums = []
    for key in keys:
        if key in key_to_num:
            nums.append(key_to_num[key])
        else:
            missing_keys.add(key)
    if not nums:
        return match.group(0)
    nums = sorted(set(nums))
    return "[" + ",".join(str(n) for n in nums) + "]"

numeric_text = re.sub(r"\[([^\[\]]*@[^]]+)\]", replace_citation, text)
OUT_MD_NUMERIC.write_text(numeric_text)

# -----------------------------
# DOCX helpers
# -----------------------------
def set_cell_text(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return run

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def add_markdown_paragraph(doc, line):
    # Bold spans **text**
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def clean_heading_text(line):
    return re.sub(r"^#+\s+", "", line).strip()

# -----------------------------
# Build DOCX
# -----------------------------
doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Times New Roman"

styles["Heading 1"].font.size = Pt(14)
styles["Heading 2"].font.size = Pt(13)
styles["Heading 3"].font.size = Pt(12)

# footer page number
footer = section.footer
footer_p = footer.paragraphs[0]
add_page_number(footer_p)

lines = numeric_text.splitlines()

for raw in lines:
    line = raw.rstrip()

    if not line.strip():
        continue

    if line.startswith("# "):
        p = doc.add_paragraph()
        p.style = doc.styles["Title"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_heading_text(line))
        run.bold = True
        continue

    if line.startswith("## "):
        heading = clean_heading_text(line)
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 1"]
        run = p.add_run(heading)
        run.bold = True
        continue

    if line.startswith("### "):
        heading = clean_heading_text(line)
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 2"]
        run = p.add_run(heading)
        run.bold = True
        continue

    if re.match(r"^\d+\.\s+", line):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        add_markdown_paragraph(doc, line)
        # Remove the blank paragraph created above if unused
        try:
            if not p.text:
                p._element.getparent().remove(p._element)
        except Exception:
            pass
        continue

    if line.startswith("- "):
        p = doc.add_paragraph(line[2:], style="List Bullet")
        continue

    add_markdown_paragraph(doc, line)

# Save
doc.save(OUT_DOCX)

# Copy into package manuscript folder
package_docx = PACKAGE_DIR / "01_manuscript" / OUT_DOCX.name
if PACKAGE_DIR.exists():
    import shutil
    shutil.copy2(OUT_DOCX, package_docx)

report = []
report.append("# Communications Biology DOCX conversion report v1\n")
report.append(f"- Input Markdown: `{IN_MD}`")
report.append(f"- Numeric-citation Markdown: `{OUT_MD_NUMERIC}`")
report.append(f"- Output DOCX: `{OUT_DOCX}`")
if PACKAGE_DIR.exists():
    report.append(f"- Copied DOCX to package: `{package_docx}`")
report.append(f"- Reference keys parsed: {len(key_to_num)}")
report.append(f"- Missing citation keys during conversion: {len(missing_keys)}")
if missing_keys:
    report.append("")
    report.append("## Missing keys")
    for key in sorted(missing_keys):
        report.append(f"- {key}")
report.append("")
report.append("## Notes")
report.append("- Citation keys were converted to numeric bracket citations using the existing reference-list order.")
report.append("- The generated DOCX is a clean Word working draft for Communications Biology formatting.")
report.append("- Final journal submission may still require reference formatting through Zotero/EndNote or journal production style.")

CONVERSION_REPORT.write_text("\n".join(report) + "\n")

print("Done.")
print(f"Wrote DOCX: {OUT_DOCX}")
print(f"Wrote numeric Markdown: {OUT_MD_NUMERIC}")
print(f"Wrote report: {CONVERSION_REPORT}")
if PACKAGE_DIR.exists():
    print(f"Copied DOCX into package: {package_docx}")
print(f"Reference keys parsed: {len(key_to_num)}")
print(f"Missing citation keys: {len(missing_keys)}")
